#!/usr/bin/env python3
"""interim-report.md 를 대회 지정 양식 docx 에 이식한다. (Python 3.10)

사용:
  <venv>/bin/python build_docx.py \
    --template "/path/중간결과보고서양식.docx" \
    --md interim-report.md --out 중간결과보고서_최종.docx

동작 규칙:
  - md 의 `## <절번호> <제목>` 이 양식의 절 헤딩 단락과 절번호로 1:1 대응한다
  - 절 헤딩 뒤 ~ 다음 절 헤딩(또는 장 헤더 표) 사이의 안내 단락·예시 표를 전부
    삭제하고 md 블록(단락/불릿/표/이미지)을 삽입한다
  - 절 헤딩·장 헤더 표 안의 "// ..." 안내문은 텍스트만 잘라낸다
  - 표지 표·날짜·지원트랙 체크를 md 의 `# 표지` 메타로 채운다
  - 삽입 런은 전부 맑은 고딕 10pt (w:eastAsia 포함)
"""
import argparse
import re
import struct
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = "맑은 고딕"
W_P = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
W_TBL = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
W_T = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"


def style_run(run, bold=False):
    run.font.name = FONT
    run.font.size = Pt(10)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)


def fill_para(par, text, bold_all=False):
    """**굵게** 인라인을 처리하며 단락에 런을 채운다."""
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if part:
            style_run(par.add_run(part), bold=bold_all or (i % 2 == 1))


def png_size(path):
    """PNG IHDR 에서 (width, height) 를 읽는다."""
    with open(path, "rb") as f:
        head = f.read(24)
    w, h = struct.unpack(">II", head[16:24])
    return w, h


def el_text(el):
    return "".join(t.text or "" for t in el.iter(W_T))


def strip_guidance(el):
    """요소 안 텍스트에서 '//' 이후를 잘라낸다 (절 헤딩·장 헤더 표용)."""
    seen = 0
    cut_at = el_text(el).find("//")
    if cut_at < 0:
        return
    for t in el.iter(W_T):
        txt = t.text or ""
        if seen >= cut_at:
            t.text = ""
        elif seen + len(txt) > cut_at:
            t.text = txt[: cut_at - seen].rstrip()
        seen += len(txt)


SEC_RE = re.compile(r"^##\s+(\d+(?:\.\d+)*)\.?\s+(.*)")
DOC_SEC_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)\.?\s+\S")


def parse_md(path):
    """반환: (표지 메타 dict, {절번호: [블록]})
    블록: ('p'|'li'|'img'|'h', str) 또는 ('tbl', [행들])"""
    cover, sections = {}, {}
    cur = None
    in_cover = False
    lines = Path(path).read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i]
        m = SEC_RE.match(ln)
        if ln.startswith("# 표지"):
            in_cover = True
        elif m:
            in_cover = False
            cur = m.group(1)
            sections[cur] = []
        elif in_cover and ln.startswith("- ") and ":" in ln:
            k, v = ln[2:].split(":", 1)
            cover[k.strip()] = v.strip()
        elif cur is not None:
            if ln.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").strip()) <= set("-: "):
                rows = []
                while i < len(lines) and lines[i].startswith("|"):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    if not set("".join(cells)) <= set("-: "):
                        rows.append(cells)
                    i += 1
                sections[cur].append(("tbl", rows))
                continue
            mm = re.match(r"^!\[.*?\]\((.+?)\)", ln)
            if mm:
                sections[cur].append(("img", mm.group(1)))
            elif ln.startswith("### "):
                sections[cur].append(("h", ln[4:]))
            elif ln.startswith("- "):
                sections[cur].append(("li", ln[2:]))
            elif ln.strip():
                sections[cur].append(("p", ln.strip()))
        i += 1
    return cover, sections


def is_chapter_table(el):
    if el.tag != W_TBL:
        return False
    first = el_text(el).strip()[:2]
    return bool(first) and first[0].isdigit() and (len(first) == 1 or not first[1].isdigit())


def build(template, mdpath, out):
    cover, sections = parse_md(mdpath)
    doc = Document(template)
    base = Path(mdpath).parent

    # 1) 표지 표 채우기
    for tbl in doc.tables:
        labels = [row.cells[0].text.strip() for row in tbl.rows]
        if "프로젝트명" not in labels:
            continue
        for row in tbl.rows:
            lab = row.cells[0].text.strip()
            if lab in cover:
                row.cells[-1].text = ""
                fill_para(row.cells[-1].paragraphs[0], cover[lab])
            elif lab == "지원트랙":
                row.cells[-1].text = ""
                fill_para(row.cells[-1].paragraphs[0], "[ ■ ] 일반 트랙  /  [   ] 국내 AI 트랙")
        break

    body = doc.element.body

    # 2) 본문 스캔: 날짜 치환, 장 헤더 표·절 헤딩의 // 안내 제거, 절 헤딩 수집
    heads = []
    for child in list(body):
        if child.tag == W_P:
            text = el_text(child).strip()
            if "2026. xx. xx" in text:
                for t in child.iter(W_T):
                    if t.text and "xx" in t.text:
                        t.text = cover.get("제출일", "2026. 07. 31") + "."
                continue
            m = DOC_SEC_RE.match(text)
            if m and not text.startswith(("//", "※", "예")):
                strip_guidance(child)
                heads.append((m.group(1), child))
        elif is_chapter_table(child):
            strip_guidance(child)

    head_els = {h[1] for h in heads}

    # 3) 절마다: 안내 구간 삭제 후 md 블록 삽입
    for num, head_el in heads:
        if num not in sections:
            continue
        sib = head_el.getnext()
        while sib is not None and sib not in head_els and not is_chapter_table(sib) and not sib.tag.endswith("sectPr"):
            nxt = sib.getnext()
            body.remove(sib)
            sib = nxt
        anchor = head_el
        for kind, data in sections[num]:
            if kind == "tbl":
                t = doc.add_table(rows=len(data), cols=len(data[0]))
                t.style = "Table Grid"
                for r, rowdata in enumerate(data):
                    for c, val in enumerate(rowdata):
                        if c >= len(t.rows[r].cells):
                            continue
                        cell = t.rows[r].cells[c]
                        cell.text = ""
                        fill_para(cell.paragraphs[0], val, bold_all=(r == 0))
                el = t._tbl
            elif kind == "img":
                p = doc.add_paragraph()
                img = base / data
                w, h = png_size(img)
                width = Cm(11) if h > w else Cm(15)
                p.add_run().add_picture(str(img), width=width)
                el = p._p
            else:
                p = doc.add_paragraph()
                prefix = "  · " if kind == "li" else ""
                fill_para(p, prefix + data, bold_all=(kind == "h"))
                el = p._p
            body.remove(el)
            anchor.addnext(el)
            anchor = el

    doc.save(out)
    print(f"저장: {out} (절 {len(sections)}개 이식)")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True)
    ap.add_argument("--md", required=True)
    ap.add_argument("--out", required=True)
    a = ap.parse_args()
    build(a.template, a.md, a.out)
